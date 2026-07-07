"""
Tools available to the agent's executor.

Each tool is a plain Python function with a predictable signature so the
executor can dispatch to it by name (this is the "tool orchestration" layer
the assignment asks about — deliberately kept as simple function dispatch
rather than a heavier framework, since the plan is already schema-constrained
JSON and doesn't need a generic tool-calling protocol on top).
"""
import os
import uuid
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)

ACCENT_COLOR = RGBColor(0x1F, 0x4E, 0x79)


# ---------------------------------------------------------------------------
# Tool 1: mock data lookup — stands in for a real system-of-record call
# (CRM, financials DB, ticketing system, etc). Swapping this for a real API
# call later means changing only this function, not the agent logic.
# ---------------------------------------------------------------------------
def mock_data_lookup(topic: str) -> dict:
    """Simulates fetching supporting data the agent doesn't have natively."""
    topic_lower = topic.lower()
    if "budget" in topic_lower or "cost" in topic_lower or "financ" in topic_lower:
        return {
            "source": "mock_finance_system",
            "estimated_budget_usd": 48000,
            "timeline_weeks": 10,
            "team_size": 4,
        }
    if "risk" in topic_lower:
        return {
            "source": "mock_risk_register",
            "top_risks": [
                "Third-party API rate limits",
                "Scope creep from stakeholder feedback",
                "Data migration downtime",
            ],
        }
    if "stakeholder" in topic_lower or "client" in topic_lower:
        return {
            "source": "mock_crm",
            "stakeholders": ["Product Owner", "Engineering Lead", "Client Sponsor"],
        }
    return {
        "source": "mock_generic_kb",
        "note": f"No specific dataset found for '{topic}'; using reasonable defaults.",
    }


# ---------------------------------------------------------------------------
# Tool 2: document generation — the mandatory final output
# ---------------------------------------------------------------------------
def generate_docx(title: str, doc_type: str, sections: list[dict], filepath: str | None = None) -> str:
    """
    Builds a polished .docx from a list of section dicts:
      {"heading": str, "body": str, "bullets": list[str] | None, "table": {"headers": [...], "rows": [[...]]} | None}
    Returns the absolute file path.
    """
    doc = Document()

    # --- Title page ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(title)
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = ACCENT_COLOR

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run(f"{doc_type} · Generated {datetime.now().strftime('%d %b %Y')}")
    sub_run.font.size = Pt(11)
    sub_run.font.italic = True

    doc.add_page_break()

    # --- Body sections ---
    for sec in sections:
        heading = doc.add_heading(sec.get("heading", "Section"), level=1)
        for r in heading.runs:
            r.font.color.rgb = ACCENT_COLOR

        body = sec.get("body")
        if body:
            doc.add_paragraph(body)

        bullets = sec.get("bullets")
        if bullets:
            for b in bullets:
                doc.add_paragraph(b, style="List Bullet")

        table_data = sec.get("table")
        if table_data and table_data.get("headers") and table_data.get("rows"):
            headers = table_data["headers"]
            rows = table_data["rows"]
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = str(h)
                for p in hdr_cells[i].paragraphs:
                    for r in p.runs:
                        r.font.bold = True
            for row in rows:
                cells = table.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = str(val)

        doc.add_paragraph()  # spacing

    if filepath is None:
        filename = f"{doc_type.lower().replace(' ', '_')}_{uuid.uuid4().hex[:8]}.docx"
        filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    return filepath


TOOL_REGISTRY = {
    "mock_data_lookup": mock_data_lookup,
    "generate_docx": generate_docx,
}
